"""Document extraction for consumer file uploads.

Matches the ChatGPT / Claude pattern: when a user uploads a file, the
orchestrator preprocesses it into LLM-ready text *before* the next
chat turn, and injects that text into ``request.metadata.attached_files``.
The agent never has to call a "read this file" tool — the extracted
content is already in its envelope.

Supported formats:

  * **PDF** — via ``pypdf`` (lazy import). Text-only; OCR-on-image-pages
    is not done (would belong in a separate OCR tool service).
  * **DOCX** — via ``python-docx`` (lazy import). Paragraphs + tables.
  * **CSV / TSV** — stdlib ``csv`` module. Header inferred; first ~100
    rows preserved with column-aligned formatting.
  * **JSON** — stdlib ``json`` module. Pretty-printed with size cap.
  * **Plain text / markdown** — passthrough with charset detection.

Hard caps:

  * Raw input ``MAX_RAW_BYTES`` (default 10 MB).
  * Extracted text ``MAX_EXTRACTED_CHARS`` (default 200_000).

Larger inputs / outputs are truncated with ``extraction_status="truncated"``.
"""
from __future__ import annotations

import csv
import io
import json
import logging
from dataclasses import dataclass, field
from typing import Any

_logger = logging.getLogger(__name__)

__all__ = [
    "ExtractedDocument",
    "ExtractionError",
    "MAX_RAW_BYTES",
    "MAX_EXTRACTED_CHARS",
    "extract_text",
    "guess_format",
]


MAX_RAW_BYTES: int = 10 * 1024 * 1024  # 10 MB
MAX_EXTRACTED_CHARS: int = 200_000


_PDF_MIMES = frozenset({"application/pdf", "application/x-pdf"})
_DOCX_MIMES = frozenset({
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
})
_CSV_MIMES = frozenset({"text/csv", "application/csv", "text/tab-separated-values"})
_JSON_MIMES = frozenset({"application/json", "text/json"})
_TEXT_MIME_PREFIX = "text/"


class ExtractionError(RuntimeError):
    """Raised when extraction fails for a reason worth surfacing."""


@dataclass(frozen=True, slots=True)
class ExtractedDocument:
    """Result of running ``extract_text`` against a blob."""

    text: str
    status: str = "ok"  # ok | truncated | unsupported | failed
    metadata: dict[str, Any] = field(default_factory=dict)


def guess_format(filename: str, content_type: str) -> str:
    """Best-effort format key from filename + content_type.

    Returns one of ``"pdf"``, ``"docx"``, ``"csv"``, ``"tsv"``, ``"json"``,
    ``"markdown"``, ``"text"``, or ``"unsupported"``. The string is
    consumed by :func:`extract_text` to pick a backend.
    """
    ct = (content_type or "").lower().split(";")[0].strip()
    fn = (filename or "").lower()
    if ct in _PDF_MIMES or fn.endswith(".pdf"):
        return "pdf"
    if ct in _DOCX_MIMES or fn.endswith(".docx"):
        return "docx"
    if fn.endswith(".tsv") or "tab-separated" in ct:
        return "tsv"
    if ct in _CSV_MIMES or fn.endswith(".csv"):
        return "csv"
    if ct in _JSON_MIMES or fn.endswith(".json"):
        return "json"
    if fn.endswith(".md") or fn.endswith(".markdown") or "markdown" in ct:
        return "markdown"
    if ct.startswith(_TEXT_MIME_PREFIX) or fn.endswith((".txt", ".log")):
        return "text"
    return "unsupported"


def _decode_text(raw: bytes) -> str:
    """Best-effort bytes → str. Tries UTF-8, chardet, then latin-1.

    UTF-16 is intentionally NOT in the fallback chain — any even-length
    byte sequence "succeeds" as UTF-16 with garbage output, which would
    silently mis-decode latin-1 / cp1252 / etc. We delegate any non-UTF-8
    case to chardet, which inspects byte patterns properly.
    """
    if not raw:
        return ""
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        pass
    # BOM-tagged UTF-16 is unambiguous and worth catching explicitly.
    if raw[:2] in (b"\xff\xfe", b"\xfe\xff"):
        try:
            return raw.decode("utf-16")
        except UnicodeDecodeError:
            pass
    try:
        import chardet  # type: ignore[import-not-found]
        detected = chardet.detect(raw[:4096])
        if detected and detected.get("encoding"):
            try:
                return raw.decode(detected["encoding"], errors="replace")
            except (UnicodeDecodeError, LookupError):
                pass
    except ImportError:
        pass
    return raw.decode("latin-1", errors="replace")


def _truncate(text: str, *, max_chars: int) -> tuple[str, bool]:
    """Return ``(text, truncated)``. Snip with an ellipsis suffix."""
    if len(text) <= max_chars:
        return text, False
    return text[: max_chars].rstrip() + "...", True


# -- Per-format extractors --------------------------------------------------


def _extract_pdf(raw: bytes, *, max_chars: int) -> ExtractedDocument:
    try:
        import pypdf  # type: ignore[import-not-found]
    except ImportError:
        return ExtractedDocument(
            text="", status="unsupported",
            metadata={"format": "pdf", "reason": "pypdf_not_installed"},
        )
    try:
        reader = pypdf.PdfReader(io.BytesIO(raw))
    except Exception as exc:  # noqa: BLE001 — pypdf raises a wide tree
        raise ExtractionError(f"pdf parse failed: {exc}") from None
    pages: list[str] = []
    n_pages = len(reader.pages)
    for idx, page in enumerate(reader.pages):
        try:
            pages.append(page.extract_text() or "")
        except Exception as exc:  # noqa: BLE001 — best-effort per page
            _logger.warning("pdf page %d extraction failed: %s", idx, exc)
            pages.append("")
    raw_text = "\n\n".join(p.strip() for p in pages if p.strip())
    text, truncated = _truncate(raw_text, max_chars=max_chars)
    return ExtractedDocument(
        text=text,
        status="truncated" if truncated else "ok",
        metadata={"format": "pdf", "n_pages": n_pages},
    )


def _extract_docx(raw: bytes, *, max_chars: int) -> ExtractedDocument:
    try:
        import docx  # type: ignore[import-not-found]
    except ImportError:
        return ExtractedDocument(
            text="", status="unsupported",
            metadata={"format": "docx", "reason": "python_docx_not_installed"},
        )
    try:
        document = docx.Document(io.BytesIO(raw))
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError(f"docx parse failed: {exc}") from None
    parts: list[str] = []
    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    n_tables = 0
    for table in document.tables:
        n_tables += 1
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    raw_text = "\n".join(parts)
    text, truncated = _truncate(raw_text, max_chars=max_chars)
    return ExtractedDocument(
        text=text,
        status="truncated" if truncated else "ok",
        metadata={
            "format": "docx",
            "n_paragraphs": len(document.paragraphs),
            "n_tables": n_tables,
        },
    )


def _extract_csv(raw: bytes, *, delimiter: str, max_chars: int) -> ExtractedDocument:
    decoded = _decode_text(raw)
    reader = csv.reader(io.StringIO(decoded), delimiter=delimiter)
    rows: list[list[str]] = []
    n_total = 0
    for row in reader:
        n_total += 1
        if len(rows) < 100:  # cap preview rows
            rows.append(row)
    if not rows:
        return ExtractedDocument(
            text="", status="ok",
            metadata={"format": "csv", "n_rows": 0, "n_columns": 0},
        )
    n_columns = max(len(r) for r in rows)
    out_lines: list[str] = []
    out_lines.append(delimiter.join(rows[0]))
    for r in rows[1:]:
        out_lines.append(delimiter.join(r))
    raw_text = "\n".join(out_lines)
    text, truncated = _truncate(raw_text, max_chars=max_chars)
    return ExtractedDocument(
        text=text,
        status="truncated" if truncated or n_total > len(rows) else "ok",
        metadata={
            "format": "tsv" if delimiter == "\t" else "csv",
            "n_rows": n_total,
            "n_preview_rows": len(rows),
            "n_columns": n_columns,
        },
    )


def _extract_json(raw: bytes, *, max_chars: int) -> ExtractedDocument:
    decoded = _decode_text(raw)
    try:
        parsed = json.loads(decoded)
    except json.JSONDecodeError as exc:
        raise ExtractionError(f"invalid json: {exc}") from None
    pretty = json.dumps(parsed, indent=2, ensure_ascii=False, default=str)
    text, truncated = _truncate(pretty, max_chars=max_chars)
    return ExtractedDocument(
        text=text,
        status="truncated" if truncated else "ok",
        metadata={"format": "json", "top_level_type": type(parsed).__name__},
    )


def _extract_text(
    raw: bytes,
    *,
    max_chars: int,
    fmt: str,
) -> ExtractedDocument:
    decoded = _decode_text(raw)
    text, truncated = _truncate(decoded, max_chars=max_chars)
    return ExtractedDocument(
        text=text,
        status="truncated" if truncated else "ok",
        metadata={"format": fmt, "n_chars": len(decoded)},
    )


# -- Public entry point ----------------------------------------------------


def extract_text(
    raw: bytes,
    *,
    filename: str = "",
    content_type: str = "",
    max_chars: int = MAX_EXTRACTED_CHARS,
    max_raw_bytes: int = MAX_RAW_BYTES,
) -> ExtractedDocument:
    """Extract LLM-ready text from a single uploaded file.

    Picks an extractor based on filename + content_type, applies hard
    size caps, and returns an :class:`ExtractedDocument`. Unsupported
    formats return ``status="unsupported"`` rather than raising —
    callers can decide whether to surface the upload as a no-text
    attachment or reject it.
    """
    if len(raw) > max_raw_bytes:
        return ExtractedDocument(
            text="",
            status="failed",
            metadata={
                "reason": "raw_size_exceeds_limit",
                "size_bytes": len(raw),
                "limit_bytes": max_raw_bytes,
            },
        )
    fmt = guess_format(filename=filename, content_type=content_type)
    try:
        if fmt == "pdf":
            return _extract_pdf(raw, max_chars=max_chars)
        if fmt == "docx":
            return _extract_docx(raw, max_chars=max_chars)
        if fmt == "csv":
            return _extract_csv(raw, delimiter=",", max_chars=max_chars)
        if fmt == "tsv":
            return _extract_csv(raw, delimiter="\t", max_chars=max_chars)
        if fmt == "json":
            return _extract_json(raw, max_chars=max_chars)
        if fmt in ("markdown", "text"):
            return _extract_text(raw, max_chars=max_chars, fmt=fmt)
    except ExtractionError as exc:
        return ExtractedDocument(
            text="", status="failed",
            metadata={"format": fmt, "reason": str(exc)},
        )
    return ExtractedDocument(
        text="", status="unsupported",
        metadata={
            "format": "unknown",
            "filename": filename,
            "content_type": content_type,
        },
    )
